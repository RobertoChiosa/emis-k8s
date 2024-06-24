from datetime import datetime

import numpy as np
from PIL import Image
from docx import Document
from fpdf import FPDF
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas


def generate_test_report_pdf() -> FPDF:
    """
    Generate a test report
    :return:
    """

    def header(document):
        """
        Set up the header of the CustomReport
        :param document: FPDF object
        :return: FPDF object
        """
        # Set up the title on the left
        document.set_font('helvetica', 'B', 15)
        document.cell(40, 10, document.title)
        # Set up the image on the right
        # document.image('../img/logo.png', 170, 6, 22)
        # Add a line break
        document.ln(5)
        # date
        document.set_font('helvetica', 'I', 8)
        document.cell(40, 10, f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        document.ln(4)
        document.cell(40, 10, f"Author: Roberto Chiosa")
        # Line break
        document.ln(10)
        document.set_line_width(0.1)
        document.line(10, 28, document.w - 10, 28)
        document.ln(5)
        return document

    def footer(document):
        """
        Set up the footer of the CustomReport
        :param document: FPDF object
        :return: FPDF object
        """
        # Position at 1.5 cm from bottom
        document.set_y(-15)
        # helvetica italic 8
        document.set_font('helvetica', 'I', 8)
        # Text color in gray
        document.set_text_color(128)
        # Page number
        # document.cell(0, 10, f"Page {document.page_no()}/{{nb}}", XPos.CENTER, new_y=YPos.NEXT)
        return document

    def add_heading(document, text, num=None, level=1):
        """
        Add a heading to the CustomReport
        :param document: FPDF object
        :param text:
        :param num:
        :param level:
        :return: FPDF object
        """
        if level == 1:
            document.set_font('helvetica', 'B', 14)
        elif level == 2:
            document.set_font('helvetica', 'B', 13)
        elif level == 3:
            document.set_font('helvetica', 'B', 12)
        # Title
        if num is None:
            document.cell(40, 10, text)
        else:
            document.cell(40, 10, f'{num}. {text}')
        # Line break
        document.ln(10)
        return document

    def add_paragraph(document, text):
        """
        Add paragraph
        :param document: FPDF object
        :param text:
        :return: FPDF object
        """
        # Read text file
        # Times 12
        document.set_font('helvetica', '', 11)
        # Output justified text
        document.multi_cell(0, 5, text)
        # Line break
        document.ln()
        return document

    def add_run(document, items_list):
        """
        :param document: FPDF object
        :param items_list:
        :return: FPDF object
        """
        # Read text file
        # Times 12
        document.set_font('helvetica', '', 11)
        for item in items_list:
            # Output justified text
            document.multi_cell(0, 5, "   - " + item)
            document.ln(2)
        document.ln()
        return document

    def add_picture(document, path):
        """
        :param document: FPDF object
        :param path:
        :return: FPDF object
        """
        # Read text file
        # Times 12
        document.image(path, x=10, y=document.get_y(), w=170)
        # Line break
        document.ln(document.get_y())
        return document

    def add_image(document, figure):
        """
        :param document: FPDF object
        :param figure:
        :return: FPDF object
        """
        canvas = FigureCanvas(figure)
        canvas.draw()
        img = Image.fromarray(np.asarray(canvas.buffer_rgba()))
        # Read text file
        # Times 12
        document.image(img, w=document.epw)
        # Line break
        # document.ln(document.get_y())
        return document

    pdf = FPDF()
    pdf.title = 'Report'

    pdf.add_page()
    pdf = header(pdf)
    pdf = add_heading(pdf, 1, 'Section')
    pdf = add_paragraph(pdf, """Section""")

    # document.add_picture('../images/fc1_definition.png')
    pdf = add_heading(pdf, 'Dataset Plot')
    pdf = add_heading(pdf, 'Dataset Statistics')
    pdf = add_run(pdf, ["- This outputs correctly", "- This outputs correctly"])
    pdf = add_heading(pdf, 'Summary Statistics filtered for when the AHU is running', level=2)
    pdf = add_heading(pdf, 'Suggestions based on data analysis', level=2)
    pdf = add_heading(pdf, 'Fault definition', level=2)
    pdf = footer(pdf)
    # document.add_table_guideline((
    #     ("Equation", "DSP < DSPSP - &DSP AND VFDSPD >= 99 % - eVFDSPD"),
    #     ("Description", "Duct static pressure too low with fan at full speed"),
    #     ("Possible Diagnosis",
    #      "- Problem with VFD\n- Mechanical problem with fan\n- Fan undersized\n- SAT set point too high (too much zone demand)")
    # ))
    return pdf


def generate_test_report_word():
    """
    Generate a test Word report
    :return:
    """
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, idx, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = idx
        row_cells[2].text = desc

    document.add_page_break()
    return document


def generate_test_report_html() -> dict:
    """
    Generate a test HTML report
    :return out: A dictionary containing the report content
    """
    out = {
        'title': 'Report',
        'subtitle': 'Report',
        'summary': {
            'title': 'Summary',
            'content': 'The aim of this report is to analyze the data quality related to a specific meter id.'
        },
        'metadata': {
            'title': 'Metadata',
            'content': 'The aim of this report is to analyze the data quality related to a specific meter id.'
        },
        'data': {
            'title': 'Data',
            'content': 'The aim of this report is to analyze the data quality related to a specific meter id.'
        },
    }
    return out
